from django.contrib.auth import login, authenticate
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib.auth import logout

from ..models import Parcel, SatisTakipModel, ParcelUserHistory
from ..serializers import ParcelSerializer
from ..filters import ParcelFilter

from rest_framework import viewsets

from django_filters.rest_framework import DjangoFilterBackend
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_protect
import json
from rest_framework.decorators import action
from rest_framework.response import Response
from django.contrib.auth.models import User


from docx import Document



@login_required
def dashboard_view(request):
    current_user = request.user

    
    return render(request, 'user_dashbord.html', {'current_user': current_user})







@csrf_protect
def approve_parcels(request):
        
    if request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))

            parcel_id = data.get("parcel_id")
            new_status = data.get("durum")

            # Parçayı veritabanından alın
            parcel = Parcel.objects.get(pk=parcel_id)

            # Yeni değerleri atayın
            parcel.durum = new_status

            # Parçayı kaydedin
            parcel.save()

            user = User.objects.get(pk=request.user.id)


            parcel_user_history = ParcelUserHistory(
                parcel= parcel,
                user=user,
                # tarih="",
                islem="onay")

            parcel_user_history.save()

            return JsonResponse({'message': 'Veriler kaydedildi.'}, status=200)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Geçersiz JSON formatı.'}, status=400)
    else:
        return JsonResponse({'error': 'Yalnızca POST istekleri kabul edilir.'}, status=405)



@csrf_protect
def remove_parcels(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))

            parcel_id = data.get("parcel_id")
            new_status = data.get("durum")
            new_user_id = data.get('user_id');
            # Parçayı veritabanından alın
            parcel = Parcel.objects.get(pk=parcel_id)

            # Yeni değerleri atayın
            parcel.durum = new_status
            parcel.user_id = new_user_id

            # Parçayı kaydedin
            parcel.save()

            user = User.objects.get(pk=request.user.id)

            parcel_user_history = ParcelUserHistory(
                parcel= parcel,
                user=user,
                # tarih="",
                islem="çıkar")

            parcel_user_history.save()

            try:
                satis_takip = SatisTakipModel.objects.get(parcel_id=parcel_id)  
                satis_takip.delete()
            except:
                pass


            return JsonResponse({'message': 'Veriler kaydedildi.'}, status=200)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Geçersiz JSON formatı.'}, status=400)
    else:
        return JsonResponse({'error': 'Yalnızca POST istekleri kabul edilir.'}, status=405)
    






@csrf_protect
def get_sales_tracking_form_data(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))

            parcel_id = data.get("parcel_id")
            print(parcel_id)
            satis_takip = SatisTakipModel.objects.get(parcel_id=parcel_id)  
                
            # SatisTakipModel verisini JSON formatında döndür
            satis_takip_data = {
                'project_name': satis_takip.project_name,
                'satis_danismani': satis_takip.satis_danismani,
                'satis_tarihi': satis_takip.satis_tarihi.strftime('%Y-%m-%d'),  # Örnek bir tarih formatı
                'referans': satis_takip.referans,
                'ada': satis_takip.ada,
                'parcel_no': satis_takip.parcel_no,
                'm2_bilgisi': satis_takip.m2_bilgisi,
                'hisse_adedi': satis_takip.hisse_adedi,
                'firma_ad_soyad': satis_takip.firma_ad_soyad,
                'TC_vergi_no': satis_takip.TC_vergi_no,
                'adres1': satis_takip.adres1,
                'adres2': satis_takip.adres2,
                'eposta': satis_takip.eposta,
                'telefon_no': satis_takip.telefon_no,
                'meslek_bilgisi': satis_takip.meslek_bilgisi,
                'ulasamadigi_durumda_aranacak_kisi': satis_takip.ulasamadigi_durumda_aranacak_kisi,
                'ulasamadigi_durumda_aranacak_kisi_telefon_no': satis_takip.ulasamadigi_durumda_aranacak_kisi_telefon_no,
                'satis_fiyati': satis_takip.satis_fiyati,
                'odeme_turu': satis_takip.odeme_turu,
                'on_odeme_tutari': satis_takip.on_odeme_tutari,
                'odeme_tarihi': satis_takip.odeme_tarihi.strftime('%Y-%m-%d'),  # Örnek bir tarih formatı
                'pesinat_tutari': satis_takip.pesinat_tutari,
                'pesinat_tarihi': satis_takip.pesinat_tarihi.strftime('%Y-%m-%d'),  # Örnek bir tarih formatı
                'm2_birim_fiyati': satis_takip.m2_birim_fiyati,
                'aciklama_bigisi': satis_takip.aciklama_bigisi,
                'ek_bilgiler': satis_takip.ek_bilgiler,
                'kampanya_kodu': satis_takip.kampanya_kodu,
                'musteri_iban': satis_takip.musteri_iban,
            }

            return JsonResponse({'data': satis_takip_data, 'message': 'Veriler alındı.'}, status=200)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Geçersiz JSON formatı.'}, status=400)
    else:
        return JsonResponse({'error': 'Yalnızca POST istekleri kabul edilir.'}, status=405)
    







@csrf_protect
def edit_form_data(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))
            parcel_id = data.get('parsel_id')
            # Veriyi kullanarak SatisTakipModel örneğini alın
            satis_takip = SatisTakipModel.objects.get(parcel_id=parcel_id)

            if data["banka_odeme"] == True:
                odeme_yontemi = "banka"
            elif data["vadeli_odeme"] == True:
                odeme_yontemi = "vadeli"
            elif data["pesin_odeme"] == True:  
                odeme_yontemi = "pesin"
            print(satis_takip.parcel_id)
            # Veriyi güncelle
            satis_takip.project_name = data.get('project_name')
            satis_takip.satis_danismani = data.get('satis_danismani')
            satis_takip.satis_tarihi = data.get('satis_tarihi')
            satis_takip.referans = data.get('referans')
            satis_takip.ada = data.get('ada')
            satis_takip.parcel_no = data.get('parcel_no')
            satis_takip.m2_bilgisi = data.get('m2_bilgisi')
            satis_takip.hisse_adedi = data.get('hisse_adedi')
            satis_takip.firma_ad_soyad = data.get('firma_ad_soyad')
            satis_takip.TC_vergi_no = data.get('TC_vergi_no')
            satis_takip.adres1 = data.get('adres1')
            satis_takip.adres2 = data.get('adres2')
            satis_takip.eposta = data.get('eposta')
            satis_takip.telefon_no = data.get('telefon_no')
            satis_takip.meslek_bilgisi = data.get('meslek_bilgisi')
            satis_takip.ulasamadigi_durumda_aranacak_kisi = data.get('ulasamadigi_durumda_aranacak_kisi')
            satis_takip.ulasamadigi_durumda_aranacak_kisi_telefon_no = data.get('ulasamadigi_durumda_aranacak_kisi_telefon_no')
            satis_takip.satis_fiyati = data.get('satis_fiyati')
            satis_takip.on_odeme_tutari = data.get('on_odeme_tutari')
            satis_takip.odeme_tarihi = data.get('odeme_tarihi')
            satis_takip.odeme_turu = odeme_yontemi
            satis_takip.pesinat_tutari = data.get('pesinat_tutari')
            satis_takip.pesinat_tarihi = data.get('pesinat_tarihi')
            satis_takip.m2_birim_fiyati = data.get('m2_birim_fiyati')
            satis_takip.aciklama_bigisi = data.get('aciklama_bigisi')
            satis_takip.ek_bilgiler = data.get('ek_bilgiler')
            satis_takip.parsel_id = data.get('parsel_id')

            # Veriyi kaydet
            satis_takip.save()

            return JsonResponse({'message': 'Veri güncellendi.'}, status=200)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Geçersiz JSON formatı.'}, status=400)
        except SatisTakipModel.DoesNotExist:
            return JsonResponse({'error': 'Belirtilen veri bulunamadı.'}, status=404)
    else:
        return JsonResponse({'error': 'Yalnızca POST istekleri kabul edilir.'}, status=405)
    













from docx.shared import Pt

def replace_text_in_cell(cell, replacements):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            for old_text, new_text in replacements.items():
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)







from django.http import FileResponse
from io import BytesIO
import io
from docx.shared import Pt  # Pt fonksiyonunu içe aktarın

@csrf_protect
def download_form_data_docx(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))

            parcel_id = data.get("parcel_id")
            print(parcel_id)
            satis_takip = SatisTakipModel.objects.get(parcel_id=parcel_id) 


            # Word belgesini açın
            belge = Document('form_template.docx')
            tablo = belge.tables[0]

            if (satis_takip.odeme_turu == "pesin"):
                pesin_ = "Peşin"
                banka_ = " "
                vadeli_ = " "
            elif (satis_takip.odeme_turu == "banka"):
                pesin_ = " "
                banka_ = "Banka"
                vadeli_ = " "
            elif (satis_takip.odeme_turu == "vadeli"):
                pesin_ = " "
                banka_ = " "
                vadeli_ = "Vadeli"

            # Değiştirme işlemleri için eski metinleri ve yeni metinleri içeren bir sözlük oluşturun
            metin_degisiklikleri = {
                'proje_ismi': satis_takip.project_name,
                'satis_danismani': satis_takip.satis_danismani,
                'satis_tarihi': satis_takip.satis_tarihi.strftime('%Y-%m-%d'),
                'referans_bilgisi': satis_takip.referans,
                'ada_bilgisi': satis_takip.ada,
                'parsel_bilgisi': satis_takip.parcel_no,
                'm2_bilgisi':  satis_takip.m2_bilgisi,
                'hisse_adedi_bilgisi': satis_takip.hisse_adedi,
                'ad_soyad_unvan_bilgisi': satis_takip.firma_ad_soyad,
                'tc_numarasi_bilgisi': satis_takip.TC_vergi_no,
                'adres1_bilgisi': satis_takip.adres1,
                'adres2_bilgisi': satis_takip.adres2,
                'eposta_bilgisi': satis_takip.eposta,
                'telefon_no_bilgisi':  satis_takip.telefon_no,
                'meslek_bilgisi':  satis_takip.meslek_bilgisi,
                'ulasilmadiginda_aranacak_kisi': satis_takip.ulasamadigi_durumda_aranacak_kisi,
                'aran_ksi_n_o_nulas':  satis_takip.ulasamadigi_durumda_aranacak_kisi_telefon_no,
                'sat_tar_bil':  satis_takip.satis_fiyati,
                'on_odem_bilgisi': satis_takip.on_odeme_tutari,
                'on_odeme_tarihi': satis_takip.odeme_tarihi.strftime('%Y-%m-%d'),
                'peşinat_tutari_bilgisi': satis_takip.pesinat_tutari,
                'pesinat_tarihi_bilgisi': satis_takip.pesinat_tarihi.strftime('%Y-%m-%d'),
                'odeme_sekli_bilgisi_pesin': pesin_,
                'odeme_sekli_bilgisi_banka': banka_,
                'odeme_sekli_bilgisi_vadeli': vadeli_,
                'm2_birim_fiyat_bilgisi': satis_takip.m2_birim_fiyati,
                'Aciklama_bilgisi': satis_takip.aciklama_bigisi,
                'ekler_bilgisi': satis_takip.ek_bilgiler,
                'musteri_banka_iban': satis_takip.musteri_iban,
            }

            # # Tablo içinde döngü ile dolaşarak metinleri değiştirin
            # for satir in tablo.rows:
            #     for hucre in satir.cells:
            #         hucre_metni = hucre.text
            #         for eski_metin, yeni_metin in metin_degisiklikleri.items():
            #             hucre_metni = hucre_metni.replace(eski_metin, yeni_metin)
            #         hucre.text = hucre_metni

            # Tablo içinde döngü ile dolaşarak metinleri değiştirin
            for satir in tablo.rows:
                for hucre in satir.cells:
                    replace_text_in_cell(hucre, metin_degisiklikleri)



            docx_bytes = io.BytesIO()
            belge.save(docx_bytes)
            docx_bytes.seek(0)

            # HttpResponse ile dosyayı kullanıcıya sunun
            response = HttpResponse(docx_bytes.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="Satış Takip Formu.docx"'
            
            return response

        except json.JSONDecodeError:
            return JsonResponse({'error': 'Geçersiz JSON formatı.'}, status=400)
        except SatisTakipModel.DoesNotExist:
            return JsonResponse({'error': 'Belirtilen veri bulunamadı.'}, status=404)
    else:
        return JsonResponse({'error': 'Yalnızca POST istekleri kabul edilir.'}, status=405)



